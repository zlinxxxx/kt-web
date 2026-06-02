"""Генератор отчёта KT5_Тестирование_TheInternet.docx по фактическим результатам."""
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

with open("results.json", "r", encoding="utf-8") as f:
    results = json.load(f)

# Метаданные тест-кейсов: предусловия, шаги, ожидаемый результат
TC_META = {
    "TC-001": {
        "name": "Авторизация с валидными данными (tomsmith)",
        "preconditions": "Страница входа /login открыта, пользователь не авторизован",
        "steps": "1. Ввести логин «tomsmith»\n2. Ввести пароль «SuperSecretPassword!»\n3. Нажать кнопку «Login»",
        "expected": "Пользователь авторизован, перенаправлен на /secure, отображается сообщение «You logged into a secure area!»",
    },
    "TC-002": {
        "name": "Авторизация с невалидным логином",
        "preconditions": "Страница входа /login открыта, сессия пользователя сброшена",
        "steps": "1. Ввести несуществующий логин «wrong_user»\n2. Ввести валидный пароль «SuperSecretPassword!»\n3. Нажать «Login»",
        "expected": "Сообщение об ошибке: «Your username is invalid!»",
    },
    "TC-003": {
        "name": "Авторизация с невалидным паролем",
        "preconditions": "Страница входа /login открыта, сессия сброшена",
        "steps": "1. Ввести логин «tomsmith»\n2. Ввести неверный пароль «wrong_password»\n3. Нажать «Login»",
        "expected": "Сообщение об ошибке: «Your password is invalid!»",
    },
    "TC-004": {
        "name": "Авторизация с пустыми полями",
        "preconditions": "Страница входа /login открыта, сессия сброшена",
        "steps": "1. Оставить поля логина и пароля пустыми\n2. Нажать «Login»",
        "expected": "Валидация сработала: отображается сообщение об ошибке",
    },
    "TC-005": {
        "name": "Переключение состояния чекбоксов",
        "preconditions": "Страница /checkboxes открыта (на ней два чекбокса)",
        "steps": "1. Зафиксировать начальное состояние чекбоксов\n2. Кликнуть на каждый чекбокс\n3. Проверить, что состояние инвертировано",
        "expected": "Состояния всех чекбоксов инвертированы по отношению к начальным",
    },
    "TC-006": {
        "name": "Выбор значения в выпадающем списке",
        "preconditions": "Страница /dropdown открыта",
        "steps": "1. Открыть выпадающий список\n2. Выбрать пункт «Option 2»\n3. Проверить текущее значение",
        "expected": "Выбранное значение — «Option 2»",
    },
    "TC-007": {
        "name": "Динамическое добавление элементов на страницу",
        "preconditions": "Страница /add_remove_elements/ открыта",
        "steps": "1. Нажать кнопку «Add Element» 3 раза\n2. Проверить количество появившихся кнопок «Delete»",
        "expected": "На странице появилось 3 кнопки «Delete» (button.added-manually)",
    },
    "TC-008": {
        "name": "Удаление динамически добавленного элемента",
        "preconditions": "На странице /add_remove_elements/ присутствует 3 кнопки «Delete»",
        "steps": "1. Нажать на первую кнопку «Delete»\n2. Проверить количество оставшихся кнопок",
        "expected": "Количество кнопок «Delete» уменьшилось на 1 (стало 2)",
    },
    "TC-009": {
        "name": "Принятие JavaScript Confirm-диалога",
        "preconditions": "Страница /javascript_alerts открыта",
        "steps": "1. Нажать кнопку «Click for JS Confirm»\n2. В появившемся диалоге нажать OK\n3. Проверить текст результата",
        "expected": "Текст alert: «I am a JS Confirm». Результат на странице: «You clicked: Ok»",
    },
    "TC-010": {
        "name": "Ожидание динамически загружаемого элемента (Dynamic Loading)",
        "preconditions": "Страница /dynamic_loading/2 открыта",
        "steps": "1. Нажать кнопку «Start»\n2. Дождаться появления элемента #finish\n3. Прочитать его текст",
        "expected": "После загрузки появляется элемент с текстом «Hello World!»",
    },
    "TC-011": {
        "name": "Появление информации о пользователе при наведении (Hovers)",
        "preconditions": "Страница /hovers открыта",
        "steps": "1. Навести курсор на первое изображение пользователя\n2. Проверить, что появилась подпись с именем",
        "expected": "Отображается подпись «name: user1»",
    },
    "TC-012": {
        "name": "Проверка корректности загрузки изображений на странице",
        "preconditions": "Страница /broken_images открыта",
        "steps": "1. Получить все элементы <img> из примера\n2. Для каждого изображения проверить naturalWidth\n3. Зафиксировать изображения с naturalWidth = 0",
        "expected": "Все изображения на странице загружены (naturalWidth > 0)",
    },
}


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tbl_pr.append(borders)


def add_kv_table(doc, rows, status_value=None):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    set_table_borders(table)
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(12)
    for i, (k, v) in enumerate(rows):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.width = Cm(4.5)
        c1.width = Cm(12)
        c0.text = ""
        c1.text = ""
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.bold = True
        r0.font.size = Pt(10)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.size = Pt(10)
        set_cell_bg(c0, "F2F2F2")
        if k == "Статус" and status_value:
            color = "C6EFCE" if status_value == "Пройден" else ("FFC7CE" if status_value == "Провален" else "FFEB9C")
            set_cell_bg(c1, color)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


doc = Document()
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# === Титул ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Отчёт по тестированию")
r.bold = True
r.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("КТ № 5. Тестирование функционала сайта")
r.bold = True
r.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("и составление тестовой документации")
r.bold = True
r.font.size = Pt(14)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Объект тестирования: тренировочный сайт The Internet (Heroku App)")
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("https://the-internet.herokuapp.com")
r.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Выполнил: Печенкин Н. А.")
r.bold = True
r.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Группа: ИТ-СИП-23-24-ДР")
r.bold = True
r.font.size = Pt(13)

doc.add_page_break()

# === 1. Описание сайта ===
add_heading(doc, "1. Описание выбранного сайта", level=1)
doc.add_paragraph("Название: The Internet — тренировочный сайт для практики UI-тестирования (автор — Дейв Хёффнер).")
doc.add_paragraph("URL: https://the-internet.herokuapp.com")
doc.add_paragraph("Тип: Демонстрационное веб-приложение со специальными страницами под отработку различных тест-сценариев.")
doc.add_paragraph(
    "The Internet — это набор изолированных страниц, каждая из которых демонстрирует отдельный аспект UI: "
    "форма авторизации, чекбоксы, выпадающие списки, JavaScript-алерты, динамически загружаемые элементы, "
    "наведение мыши, битые изображения, статус-коды и т. д. Сайт специально создан для отработки навыков "
    "автоматизированного тестирования и содержит ряд намеренно встроенных «дефектов» (например, страница "
    "/broken_images), что позволяет проверить способность тестировщика их обнаружить."
)

add_heading(doc, "Обоснование выбора", level=2)
for item in [
    "Сайт стабильно доступен и предназначен именно для практики автоматизации",
    "Содержит широкий спектр UI-элементов: формы, alert'ы, динамический контент, hover-эффекты",
    "Возможность автоматизированного тестирования через Selenium WebDriver",
    "Присутствуют намеренные дефекты, позволяющие отработать процесс баг-репортинга",
]:
    doc.add_paragraph(item, style="List Bullet")

# === 2. Тест-план ===
add_heading(doc, "2. Тестовый план", level=1)

add_heading(doc, "2.1 Объём тестирования", level=2)
for item in [
    "Авторизация (валидные данные, невалидный логин, невалидный пароль, пустые поля)",
    "Работа с интерактивными элементами (чекбоксы, выпадающий список)",
    "Динамическое добавление и удаление элементов на странице",
    "Обработка JavaScript-диалогов (Confirm)",
    "Ожидание динамически загружаемых элементов",
    "Реакция UI на наведение курсора (hover)",
    "Проверка корректности загрузки изображений",
]:
    doc.add_paragraph(item, style="List Bullet")

add_heading(doc, "2.2 Подход к тестированию", level=2)
doc.add_paragraph(
    "Автоматизированное тестирование с использованием Python + Selenium WebDriver. Техника чёрного ящика: "
    "тестирование на основе ожидаемого поведения без анализа исходного кода. Применены методики "
    "эквивалентного разделения (валидные/невалидные данные), негативного тестирования (пустые поля, "
    "неверные учётные данные) и проверки фактического состояния DOM."
)

add_heading(doc, "2.3 Критерии начала тестирования", level=2)
for item in [
    "Сайт the-internet.herokuapp.com доступен",
    "Selenium WebDriver и ChromeDriver установлены",
    "Тестовые сценарии разработаны и оформлены в виде скрипта",
]:
    doc.add_paragraph(item, style="List Bullet")

add_heading(doc, "2.4 Критерии завершения тестирования", level=2)
for item in [
    "Все 12 тестовых сценариев выполнены",
    "Все найденные дефекты зафиксированы",
    "Отчёт оформлен",
]:
    doc.add_paragraph(item, style="List Bullet")

add_heading(doc, "2.5 Среда тестирования", level=2)
add_kv_table(
    doc,
    [
        ("Параметр", "Значение"),
        ("ОС", "Windows 11 (PyCharm IDE)"),
        ("Браузер", "Google Chrome (через Selenium WebDriver)"),
        ("Язык автоматизации", "Python 3 + Selenium + webdriver-manager"),
        ("Разрешение", "1920 × 1080"),
        ("Тип тестирования", "Автоматизированное функциональное (чёрный ящик)"),
        ("Дата тестирования", "28 мая 2026"),
    ],
)

# === 3. Тест-кейсы ===
add_heading(doc, "3. Тестовые сценарии и тест-кейсы", level=1)
for r in results:
    meta = TC_META[r["id"]]
    rows = [
        ("Поле", "Значение"),
        ("ID", r["id"]),
        ("Название", meta["name"]),
        ("Предусловия", meta["preconditions"]),
        ("Шаги", meta["steps"]),
        ("Ожидаемый результат", meta["expected"]),
        ("Фактический результат", r["actual"]),
        ("Статус", r["status"]),
    ]
    add_kv_table(doc, rows, status_value=r["status"])

# === 4. Результаты ===
add_heading(doc, "4. Результаты тестирования", level=1)
add_heading(doc, "4.1 Сводка", level=2)
passed = sum(1 for x in results if x["status"] == "Пройден")
failed = sum(1 for x in results if x["status"] == "Провален")
errors = sum(1 for x in results if x["status"] == "Ошибка")
total = len(results)
defects_count = failed + errors
percent = round(passed / total * 100)
add_kv_table(
    doc,
    [
        ("Показатель", "Значение"),
        ("Всего тест-кейсов", str(total)),
        ("Пройдено", str(passed)),
        ("Провалено", str(failed + errors)),
        ("Процент успешности", f"{percent}%"),
        ("Найдено дефектов", str(defects_count)),
    ],
)

# === 4.2 Дефекты ===
add_heading(doc, "4.2 Найденные дефекты", level=2)

defects = []
bug_id = 1
for r in results:
    if r["status"] in ("Провален", "Ошибка"):
        if r["id"] == "TC-012":
            defects.append(
                {
                    "id": f"BUG-{bug_id:03d}",
                    "name": "Битые изображения на странице /broken_images",
                    "tc": r["id"],
                    "description": (
                        "На странице https://the-internet.herokuapp.com/broken_images среди трёх "
                        "изображений два не загружаются: asdf.jpg и hjkl.jpg (HTTP 404, naturalWidth = 0). "
                        "Пользователю отображается стандартная иконка «битой картинки». Это нарушает "
                        "восприятие контента и снижает доверие к сайту."
                    ),
                    "severity": "Средняя",
                    "priority": "Средний",
                    "recommendation": (
                        "Заменить пути к изображениям asdf.jpg и hjkl.jpg на корректные URL "
                        "существующих файлов. Дополнительно — реализовать обработчик onerror "
                        "у тегов <img>, чтобы при ошибке загрузки отображалась заглушка."
                    ),
                }
            )
        else:
            defects.append(
                {
                    "id": f"BUG-{bug_id:03d}",
                    "name": f"Тест-кейс {r['id']} не прошёл",
                    "tc": r["id"],
                    "description": r["actual"],
                    "severity": "Средняя",
                    "priority": "Средний",
                    "recommendation": "Проверить ожидаемое поведение функционала и исправить расхождение.",
                }
            )
        bug_id += 1

if not defects:
    doc.add_paragraph("В ходе автоматизированного тестирования дефектов не обнаружено.")
else:
    for d in defects:
        rows = [
            ("Поле", "Значение"),
            ("ID дефекта", d["id"]),
            ("Название", d["name"]),
            ("Связанный тест-кейс", d["tc"]),
            ("Описание", d["description"]),
            ("Серьёзность", d["severity"]),
            ("Приоритет", d["priority"]),
            ("Рекомендация", d["recommendation"]),
        ]
        add_kv_table(doc, rows)

# === 5. Рекомендации ===
add_heading(doc, "5. Рекомендации", level=1)
if defects:
    doc.add_paragraph(
        f"По результатам автоматизированного тестирования сайта the-internet.herokuapp.com "
        f"выявлено {len(defects)} дефект(а):"
    )
    for d in defects:
        doc.add_paragraph(d["recommendation"], style="List Number")
    doc.add_paragraph("Провести регрессионное тестирование после устранения дефектов.", style="List Number")
else:
    doc.add_paragraph(
        "Все тестируемые сценарии отработали в соответствии с ожидаемым поведением. "
        "Рекомендуется расширить покрытие тестами форм с загрузкой файлов и страниц с "
        "iframe / окнами авторизации (Basic Auth)."
    )

# === 6. Заключение ===
add_heading(doc, "6. Заключение", level=1)
conclusion = (
    f"В ходе автоматизированного тестирования веб-приложения the-internet.herokuapp.com "
    f"с использованием Python и Selenium WebDriver было выполнено {total} тест-кейсов, "
    f"охватывающих авторизацию (4 сценария), работу с интерактивными элементами (4 сценария: "
    f"чекбоксы, выпадающий список, добавление и удаление элементов), JavaScript-диалоги, "
    f"динамическую загрузку, hover-эффекты и проверку изображений. "
    f"{passed} из {total} тестов пройдены успешно ({percent}%). "
)
if defects:
    conclusion += (
        f"Обнаружено {len(defects)} дефект(а) средней серьёзности, все зафиксированы "
        f"с подробным описанием и рекомендациями по исправлению. "
    )
else:
    conclusion += "Дефектов не обнаружено. "
conclusion += "Тестирование проведено 28 мая 2026 года."
doc.add_paragraph(conclusion)

out = "KT5_Тестирование_TheInternet.docx"
doc.save(out)
print(f"Отчёт сохранён: {out}")
